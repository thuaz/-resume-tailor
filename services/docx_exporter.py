"""Generate a clean .docx file from the Claude-tailored resume text.

Parses the structured output format (## NAME, ## CONTACT, ...) and renders
each section with appropriate Word styles. No watermarks, no headers/footers.
"""
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import EXPORT_DIR


# ── Public API ─────────────────────────────────────────────────────────────


def export_to_docx(
    tailored_text: str, company_hint: str = "", role_hint: str = ""
) -> Path:
    """Convert a Claude-tailored resume to a formatted .docx file.

    Args:
        tailored_text: The full output from Claude (with ## SECTION markers).
        company_hint: Optional company name for the filename.
        role_hint: Optional role name for the filename.

    Returns:
        Path to the saved .docx file.
    """
    doc = _build_document(tailored_text)
    path = _make_output_path(company_hint, role_hint)
    doc.save(str(path))
    return path


# ── Document construction ──────────────────────────────────────────────────


def _build_document(text: str) -> Document:
    doc = Document()

    _setup_page(doc)
    _setup_default_style(doc)

    sections = _parse_sections(text)

    for sec_type, sec_content in sections:
        if sec_type == "NAME":
            _render_name(doc, sec_content)
        elif sec_type == "CONTACT":
            _render_contact(doc, sec_content)
        elif sec_type in ("PROFESSIONAL SUMMARY", "SUMMARY"):
            _render_section_heading(doc, "PROFESSIONAL SUMMARY")
            _render_body(doc, sec_content)
        elif sec_type == "SKILLS":
            _render_section_heading(doc, "SKILLS")
            _render_skills(doc, sec_content)
        elif sec_type == "EXPERIENCE":
            _render_section_heading(doc, "EXPERIENCE")
            _render_experience(doc, sec_content)
        elif sec_type == "EDUCATION":
            _render_section_heading(doc, "EDUCATION")
            _render_body(doc, sec_content)
        elif sec_type == "PROJECTS":
            _render_section_heading(doc, "PROJECTS")
            _render_experience(doc, sec_content)
        elif sec_type == "RECOMMENDATIONS":
            _render_section_heading(doc, "RECOMMENDATIONS")
            _render_body(doc, sec_content)
        else:
            # Unknown section — render as plain body
            if sec_content.strip():
                _render_section_heading(doc, sec_type.upper())
                _render_body(doc, sec_content)

    return doc


# ── Page & style setup ─────────────────────────────────────────────────────


def _setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def _setup_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15


# ── Section renderers ──────────────────────────────────────────────────────


def _render_name(doc: Document, content: str) -> None:
    name = content.strip().split("\n")[0].strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.size = Pt(18)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)


def _render_contact(doc: Document, content: str) -> None:
    contact_line = content.strip().replace("\n", "  |  ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(contact_line)
    run.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(2)

    # Thin separator line after contact info
    _add_horizontal_rule(doc)


def _render_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.size = Pt(12)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

    # Bottom border on the paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "333333",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_body(doc: Document, content: str) -> None:
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("- "):
            _render_bullet(doc, line[2:])
        elif line.startswith("* "):
            _render_bullet(doc, line[2:])
        else:
            p = doc.add_paragraph(line)
            # shrink space for compact body
            p.paragraph_format.space_after = Pt(1)


def _render_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.strip(), style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)


def _render_skills(doc: Document, content: str) -> None:
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            cat, skills = line.split(":", 1)
            p = doc.add_paragraph()
            run_cat = p.add_run(cat.strip() + ": ")
            run_cat.bold = True
            run_cat.size = Pt(10.5)
            run_cat.font.name = "Calibri"
            run_skills = p.add_run(skills.strip())
            run_skills.size = Pt(10.5)
            run_skills.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)
        else:
            _render_body(doc, line)


def _render_experience(doc: Document, content: str) -> None:
    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            header = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.size = Pt(10.5)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
        elif line.startswith("- ") or line.startswith("* "):
            _render_bullet(doc, line[2:])
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(1)


# ── Helpers ────────────────────────────────────────────────────────────────


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin line by putting a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_sections(text: str) -> list[tuple[str, str]]:
    """Split Claude output into (section_name, section_content) pairs.

    Section headers are lines matching '## NAME', '## CONTACT', etc.
    """
    pattern = r"^## (.+)$"
    parts = re.split(pattern, text, flags=re.MULTILINE)
    # parts[0] is any text before the first ## heading (discard)
    # Then alternating: heading_name, content
    sections: list[tuple[str, str]] = []
    for i in range(1, len(parts), 2):
        name = parts[i].strip()
        content = parts[i + 1].strip() if i + 1 < len(parts) else ""
        sections.append((name, content))
    return sections


def _make_output_path(company_hint: str = "", role_hint: str = "") -> Path:
    """Generate a clean filename and return the full output path."""
    parts = ["resume"]
    if company_hint:
        parts.append(re.sub(r"[^a-zA-Z0-9_-]", "", company_hint))
    if role_hint:
        parts.append(re.sub(r"[^a-zA-Z0-9_-]", "", role_hint))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    parts.append(timestamp)
    filename = "_".join(parts) + ".docx"
    return EXPORT_DIR / filename
