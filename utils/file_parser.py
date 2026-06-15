"""Extract plain text from uploaded resume files (.docx and .pdf)."""
import io

from docx import Document


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract all paragraph text from a .docx file."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also check tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract all page text from a .pdf file using pdfplumber."""
    import pdfplumber

    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return "\n\n".join(pages)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect format and extract text.

    Returns the extracted plain text, or raises ValueError for unsupported types.
    """
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_from_docx(file_bytes)
    elif lower.endswith(".pdf"):
        return extract_from_pdf(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {filename}")
